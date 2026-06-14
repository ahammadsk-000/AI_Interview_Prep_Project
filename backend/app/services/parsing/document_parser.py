"""Document text extraction for PDF / DOCX / TXT, with secure-upload validation.

Each parser is small and isolated (SRP). ``extract_text`` dispatches by detected
file type. Validation rejects unknown extensions, oversized files, and empty
uploads before any parsing touches the bytes.
"""
from __future__ import annotations

import io
import os
import re
import unicodedata

from app.core.config import settings
from app.core.exceptions import ValidationError
from app.domain.resume.enums import FileType

_EXT_TO_TYPE = {".pdf": FileType.PDF, ".docx": FileType.DOCX, ".txt": FileType.TXT}
_EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}
# Magic-byte signatures for defense-in-depth (extension can be spoofed).
_PDF_MAGIC = b"%PDF-"
_ZIP_MAGIC = b"PK\x03\x04"  # docx is a zip container


def safe_filename(filename: str) -> str:
    """Strip path components and unsafe characters from a client-supplied name."""
    base = os.path.basename(filename or "").strip()
    base = unicodedata.normalize("NFKD", base)
    base = re.sub(r"[^A-Za-z0-9._-]", "_", base)
    return base[:200] or "upload"


def detect_type(filename: str) -> FileType:
    ext = os.path.splitext(filename.lower())[1]
    if ext not in _EXT_TO_TYPE:
        raise ValidationError(
            f"Unsupported file type '{ext or 'unknown'}'. "
            f"Allowed: {', '.join(settings.ALLOWED_UPLOAD_EXTENSIONS)}."
        )
    return _EXT_TO_TYPE[ext]


def mime_for(filename: str) -> str:
    return _EXT_TO_MIME[os.path.splitext(filename.lower())[1]]


def validate_upload(filename: str, content: bytes) -> FileType:
    if not content:
        raise ValidationError("Uploaded file is empty.")
    if len(content) > settings.MAX_UPLOAD_BYTES:
        mb = settings.MAX_UPLOAD_BYTES // (1024 * 1024)
        raise ValidationError(f"File exceeds the {mb} MB upload limit.")
    file_type = detect_type(filename)
    # Magic-byte sanity check.
    if file_type is FileType.PDF and not content.startswith(_PDF_MAGIC):
        raise ValidationError("File extension is .pdf but content is not a valid PDF.")
    if file_type is FileType.DOCX and not content.startswith(_ZIP_MAGIC):
        raise ValidationError("File extension is .docx but content is not a valid DOCX.")
    return file_type


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _parse_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def extract_text(content: bytes, file_type: FileType) -> str:
    parsers = {
        FileType.PDF: _parse_pdf,
        FileType.DOCX: _parse_docx,
        FileType.TXT: _parse_txt,
    }
    try:
        text = parsers[file_type](content)
    except ValidationError:
        raise
    except Exception as exc:  # pragma: no cover - corrupt file path
        raise ValidationError(f"Could not extract text from the document: {exc}") from exc
    return text.strip()
